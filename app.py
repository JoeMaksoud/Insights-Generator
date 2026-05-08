import streamlit as st
import os
import io
import json
import re
from datetime import datetime

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="EXD Insights Generator",
    page_icon="⚡",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {
  font-family: 'Inter', sans-serif !important;
  background-color: #0f0f0f !important;
  color: #d8d8d8 !important;
}

#MainMenu, footer, header, .stDeployButton { visibility: hidden !important; }

.block-container {
  max-width: 800px !important;
  padding: 2rem 2rem 5rem !important;
}

/* ── Section headers ── */
.section-hdr {
  display: flex;
  align-items: center;
  gap: 12px;
  margin: 2.5rem 0 1rem;
}
.section-num {
  min-width: 28px; height: 28px;
  border-radius: 50%;
  background: #2a2a2a;
  color: #fff;
  font-size: 13px; font-weight: 600;
  display: flex; align-items: center; justify-content: center;
  flex-shrink: 0;
}
.section-ttl { font-size: 15px; font-weight: 600; color: #fff; }
.section-sub { font-size: 13px; color: #555; font-weight: 400; }

/* ── Inputs ── */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
  background: #1a1a1a !important;
  border: 1px solid #2a2a2a !important;
  border-radius: 6px !important;
  color: #e0e0e0 !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 14px !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
  border-color: #ff6b2b !important;
  box-shadow: 0 0 0 2px #ff6b2b18 !important;
}
.stTextInput > label, .stTextArea > label,
.stSelectbox > label, .stMultiSelect > label {
  font-size: 12px !important; font-weight: 600 !important;
  letter-spacing: 0.05em !important; color: #555 !important;
  text-transform: uppercase !important;
}

/* ── Selectbox ── */
.stSelectbox [data-baseweb="select"] > div {
  background: #1a1a1a !important;
  border: 1px solid #2a2a2a !important;
  border-radius: 6px !important;
  color: #e0e0e0 !important;
  font-family: 'Inter', sans-serif !important;
  font-size: 14px !important;
}
.stSelectbox [data-baseweb="select"] > div:focus-within {
  border-color: #ff6b2b !important;
}

/* ── ALL buttons: ghost dark by default ── */
div[data-testid="stButton"] > button {
  background: #1e1e1e !important;
  color: #ccc !important;
  border: 1px solid #2a2a2a !important;
  border-radius: 6px !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 600 !important;
  font-size: 13px !important;
  padding: 0.5rem 1rem !important;
  transition: all 0.15s !important;
  width: 100% !important;
}
div[data-testid="stButton"] > button:hover {
  border-color: #ff6b2b !important;
  color: #ff6b2b !important;
  background: #ff6b2b0f !important;
}
div[data-testid="stButton"] > button:disabled {
  opacity: 0.3 !important;
  cursor: not-allowed !important;
}

/* ── Generate button — full width, dark fill ── */
div[data-testid="stButton"].generate-btn > button,
div[data-testid="stButton"]:has(button[kind="primary"]) > button {
  background: #1a1a1a !important;
  color: #fff !important;
  border: 1px solid #333 !important;
  font-size: 15px !important;
  padding: 0.75rem 2rem !important;
  letter-spacing: 0.03em !important;
}

/* ── Download buttons ── */
.stDownloadButton > button {
  background: #1e1e1e !important;
  color: #ccc !important;
  border: 1px solid #2a2a2a !important;
  border-radius: 6px !important;
  font-family: 'Inter', sans-serif !important;
  font-weight: 600 !important;
  font-size: 13px !important;
  padding: 0.5rem 1rem !important;
  transition: all 0.15s !important;
}
.stDownloadButton > button:hover {
  border-color: #ff6b2b !important;
  color: #ff6b2b !important;
}

/* ── Spinner ── */
.stSpinner > div { border-top-color: #ff6b2b !important; }

/* ── Divider ── */
hr { border-color: #1e1e1e !important; margin: 1.5rem 0 !important; }

/* ── Expander ── */
.streamlit-expanderHeader {
  background: #1a1a1a !important;
  border-radius: 6px !important;
  font-size: 13px !important;
  color: #777 !important;
}

/* ── Vertical cards ── */
.vert-card {
  background: #1a1a1a;
  border: 1px solid #2a2a2a;
  border-radius: 8px;
  padding: 0.9rem 0.75rem 0.5rem;
  text-align: center;
  margin-bottom: 0.4rem;
}
.vert-card.selected {
  border-color: #ff6b2b;
  background: #ff6b2b0d;
}
.vert-icon { font-size: 22px; display: block; margin-bottom: 0.35rem; }
.vert-name {
  font-size: 12px; font-weight: 600; color: #ccc;
  line-height: 1.3; display: block;
}
.vert-card.selected .vert-name { color: #ff6b2b; }

/* ── Output option cards ── */
.opt-card {
  background: #1a1a1a;
  border: 1px solid #2a2a2a;
  border-radius: 8px;
  padding: 0.85rem 1rem;
  margin-bottom: 0.5rem;
  cursor: pointer;
  display: flex;
  align-items: center;
  gap: 0.75rem;
}
.opt-card.selected {
  border-color: #ff6b2b;
  background: #ff6b2b0d;
}
.opt-card-title { font-size: 14px; font-weight: 600; color: #e0e0e0; }
.opt-card-sub { font-size: 12px; color: #555; margin-top: 2px; }
.opt-dot {
  width: 16px; height: 16px; border-radius: 50%;
  border: 2px solid #333; flex-shrink: 0;
  display: flex; align-items: center; justify-content: center;
}
.opt-dot.selected {
  border-color: #ff6b2b; background: #ff6b2b;
}
.opt-dot.selected::after {
  content: ''; width: 6px; height: 6px;
  border-radius: 50%; background: #fff;
}

/* ── Result cards ── */
.res-card {
  background: #141414;
  border: 1px solid #202020;
  border-left: 3px solid #ff6b2b;
  border-radius: 8px;
  padding: 1rem 1.2rem;
  margin-bottom: 0.65rem;
  font-size: 13.5px; color: #c8c8c8; line-height: 1.7;
}
.stat-card {
  background: #141414;
  border: 1px solid #202020;
  border-radius: 8px;
  padding: 1rem 1.2rem;
  margin-bottom: 0.6rem;
}
.stat-num { font-size: 1rem; font-weight: 700; color: #ff6b2b; line-height: 1.3; margin-bottom: 0.15rem; }
.stat-src { font-size: 11px; color: #333; font-family: 'Courier New', monospace; margin-bottom: 0.35rem; }
.stat-ctx { font-size: 12.5px; color: #777; line-height: 1.5; }
.rec-card {
  background: #ff6b2b0a;
  border: 1px solid #ff6b2b22;
  border-radius: 8px;
  padding: 1rem 1.2rem;
  margin-top: 0.65rem;
}
.rec-label {
  font-size: 10px; font-weight: 700;
  letter-spacing: 0.15em; text-transform: uppercase;
  color: #ff6b2b; margin-bottom: 0.4rem;
}
.rec-text { font-size: 13px; color: #ccc; line-height: 1.65; }
.headline-pill {
  font-size: 14px; font-weight: 600; color: #e8e8e8;
  border-left: 3px solid #ff6b2b;
  padding: 0.4rem 0 0.4rem 0.85rem;
  margin-bottom: 1rem; line-height: 1.45;
}
.vert-badge {
  display: inline-flex; align-items: center; gap: 0.4rem;
  font-size: 11px; font-weight: 700;
  letter-spacing: 0.1em; text-transform: uppercase;
  color: #ff6b2b;
  background: #ff6b2b12;
  border: 1px solid #ff6b2b30;
  border-radius: 4px;
  padding: 0.2rem 0.6rem; margin-bottom: 0.9rem;
}
.result-header {
  padding: 1.25rem 0 0.5rem;
  border-bottom: 1px solid #1e1e1e;
  margin-bottom: 1.5rem;
}
.result-header h2 {
  font-size: 1.4rem !important; font-weight: 700 !important;
  color: #f0f0f0 !important; margin: 0.3rem 0 0.2rem !important;
}
.result-header p { font-size: 12px; color: #444; margin: 0; }
.step-row {
  display: flex; align-items: center; gap: 8px;
  font-size: 12px; color: #444;
  font-family: 'Courier New', monospace; padding: 0.3rem 0;
}
.step-dot { width: 6px; height: 6px; background: #ff6b2b; border-radius: 50%; flex-shrink: 0; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# GEMINI
# ─────────────────────────────────────────────────────────────────────────────

@st.cache_resource
def get_gemini_client():
    try:
        import google.genai as genai
    except ImportError:
        st.error("google-genai not installed.")
        st.stop()
    api_key = None
    try:
        api_key = st.secrets["GEMINI_KEY"]
    except Exception:
        api_key = os.environ.get("GEMINI_KEY")
    if not api_key:
        st.error("⚠️ GEMINI_KEY not found in Streamlit secrets.")
        st.stop()
    return genai.Client(api_key=api_key)


def call_gemini(client, prompt: str, grounded: bool = True) -> str:
    from google.genai import types

    # Use gemini-2.5-flash (stable, available to all keys)
    MODEL = "gemini-2.5-flash-preview-04-17"

    def _call(use_grounding: bool) -> str:
        kwargs = {"temperature": 0.7, "max_output_tokens": 4096}
        if use_grounding:
            kwargs["tools"] = [types.Tool(google_search=types.GoogleSearch())]
        response = client.models.generate_content(
            model=MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(**kwargs),
        )
        if response.text:
            return response.text
        parts = []
        for candidate in response.candidates:
            for part in candidate.content.parts:
                if hasattr(part, "text") and part.text:
                    parts.append(part.text)
        return "\n".join(parts)

    if grounded:
        try:
            return _call(True)
        except Exception as e:
            err = str(e).lower()
            if any(x in err for x in ["grounding", "search", "tool", "400", "403", "permission", "not_found", "404"]):
                st.warning("Google Search grounding unavailable — generating without live web data.", icon="🔍")
                return _call(False)
            raise
    return _call(False)


def parse_json(text: str):
    if not text:
        return None
    match = re.search(r"```(?:json)?\s*([\[\{].*?)\s*```", text, re.DOTALL)
    raw = match.group(1) if match else text.strip()
    start = min(
        (raw.find("[") if raw.find("[") != -1 else len(raw)),
        (raw.find("{") if raw.find("{") != -1 else len(raw)),
    )
    end = max(raw.rfind("}"), raw.rfind("]"))
    if 0 <= start <= end:
        raw = raw[start:end + 1]
    try:
        return json.loads(raw)
    except Exception:
        return None


# ─────────────────────────────────────────────────────────────────────────────
# DATA
# ─────────────────────────────────────────────────────────────────────────────

MARKETS = [
    "UAE", "Saudi Arabia", "Qatar", "Kuwait", "Bahrain", "Oman",
    "Egypt", "Jordan", "Lebanon", "Morocco", "South Africa",
    "United Kingdom", "United States", "France", "Germany",
    "India", "China", "Japan", "Australia",
]

VERTICALS = {
    "SEO & AI Search": {
        "icon": "🔍",
        "focus": (
            "organic search, AI search engines (Perplexity, ChatGPT Search, Google AI Overviews), "
            "GEO (Generative Engine Optimization), content structure for AI citation, entity authority, "
            "E-E-A-T, zero-click trends, SGE/AIO impact on category traffic, crawlability"
        ),
    },
    "Customer Experience": {
        "icon": "💬",
        "focus": (
            "CX strategy, personalisation at scale, loyalty programmes, omnichannel journeys, "
            "voice of customer, NPS/CSAT benchmarks, digital-to-physical integration, "
            "conversational AI, CRM maturity, retention economics"
        ),
    },
    "Tech & Creative": {
        "icon": "⚙️",
        "focus": (
            "martech stack, AI-generated creative at scale, personalisation engines, A/B testing, "
            "Core Web Vitals, accessibility, design systems, headless CMS, "
            "composable architecture, creative production velocity"
        ),
    },
    "Strategy & Growth": {
        "icon": "📈",
        "focus": (
            "paid media efficiency, attribution, growth loops, market share dynamics, "
            "competitive positioning, share of search, category entry points, "
            "commercial strategy, pricing perception, media mix modelling"
        ),
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# PROMPTS
# ─────────────────────────────────────────────────────────────────────────────

def market_stats_prompt(client_name, market, industry, website, fmt):
    if fmt == "headline":
        output_fmt = (
            "Return exactly 6 statistics as a JSON array. Each item: "
            '{ "stat": "the number/finding (max 15 words)", '
            '"source": "Publication Name, Year", '
            '"context": "one sentence of relevance (max 20 words)" }. '
            "Return ONLY the JSON array."
        )
    else:
        output_fmt = (
            "Return a 220-word narrative paragraph with at least 6 inline statistics "
            "and citations (e.g. 'according to Bain & Co, 2024'). "
            'Return as JSON: { "narrative": "..." }. Return ONLY the JSON.'
        )
    return f"""You are a senior market intelligence analyst preparing a new business pitch.

CLIENT: {client_name}
MARKET: {market}
INDUSTRY: {industry}
WEBSITE: {website}

Search for current, specific, and surprising statistics about this market and industry.
- Use real data from credible sources (Euromonitor, Statista, analyst firms, government bodies, news)
- Be specific to {market} — not generic global figures unless no local data exists
- Prioritise 2023–2025 data
- Cover: market size/growth, digital adoption, consumer behaviour, AI/tech disruption, competitive dynamics
- Do not fabricate. Use directional language if exact figures unavailable.

{output_fmt}"""


def vertical_prompt(client_name, market, industry, website, vertical, focus, include_narrative):
    narrative_key = (
        '\n- "narrative": 120-word strategic paragraph connecting insights into a pitch story'
        if include_narrative else ""
    )
    return f"""You are EXD — the Experience Design practice inside Performics / Publicis Groupe — preparing a new business pitch.

CLIENT: {client_name} | MARKET: {market} | INDUSTRY: {industry} | WEBSITE: {website}
VERTICAL: {vertical} | FOCUS: {focus}

Search for real, current intelligence about this client's digital presence and competitive landscape.

Return a JSON object with:
- "headline": punchy specific headline (max 12 words). Must name the client or market — never generic.
- "insights": array of 5 bullets. Each must:
  * Be specific to {market}/{industry} — not a recycled generic claim
  * Include a real data point, competitor name, platform, or named trend
  * First sentence = the fact. Second sentence = the implication.
  * 2–3 sentences total.
- "recommendation": bold EXD recommendation (2–3 sentences). Start with an action verb. Specific to {vertical}.{narrative_key}

QUALITY BAR:
BAD: "74% of users expect personalised experiences."
GOOD: "Emirates.com returns no AI Overview for 23 of the top 30 aviation queries in the UAE — a gap Etihad is actively closing by restructuring FAQ content for Perplexity citations."

Return ONLY the JSON object. No preamble."""


# ─────────────────────────────────────────────────────────────────────────────
# WORD EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(client_name, market, industry, results) -> bytes | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
    except ImportError:
        return None

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    ORANGE = RGBColor(255, 107, 43)
    WHITE  = RGBColor(240, 240, 240)
    GREY   = RGBColor(140, 140, 140)
    DIM    = RGBColor(80, 80, 80)

    def add(text="", bold=False, size=10, color=None, italic=False, after=6):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(after)
        if text:
            run = para.add_run(text)
            run.font.name = "Arial"; run.font.size = Pt(size)
            run.font.bold = bold; run.font.italic = italic
            run.font.color.rgb = color or WHITE
        return para

    def rule():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("─" * 88)
        r.font.name = "Arial"; r.font.size = Pt(7)
        r.font.color.rgb = RGBColor(35, 35, 35)

    add("PERFORMICS · EXD PRACTICE", bold=True, size=7, color=ORANGE, after=2)
    add("INSIGHTS BRIEF", bold=True, size=20, color=WHITE, after=4)
    add(f"{client_name}  ·  {market}  ·  {industry}", size=10, color=GREY, after=2)
    add(f"Generated {datetime.now().strftime('%d %B %Y')}", size=8, color=DIM, after=12)
    rule()

    if "stats" in results:
        add("MARKET INTELLIGENCE", bold=True, size=8, color=ORANGE, after=6)
        stats = results["stats"]
        if isinstance(stats, list):
            for s in stats:
                if isinstance(s, dict):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    r1 = p.add_run(f"{s.get('stat','')}  ")
                    r1.font.name = "Arial"; r1.font.size = Pt(10)
                    r1.font.bold = True; r1.font.color.rgb = ORANGE
                    r2 = p.add_run(f"({s.get('source','')})")
                    r2.font.name = "Arial"; r2.font.size = Pt(8); r2.font.color.rgb = DIM
                    if s.get("context"):
                        add(s["context"], size=9, color=GREY, after=7)
        elif isinstance(stats, dict) and "narrative" in stats:
            add(stats["narrative"], size=9, color=GREY, after=8)
        rule()

    for vertical, data in results.get("verticals", {}).items():
        icon = VERTICALS.get(vertical, {}).get("icon", "")
        add(f"{icon}  {vertical.upper()}", bold=True, size=8, color=ORANGE, after=4)
        if data.get("headline"):
            add(data["headline"], bold=True, size=12, color=WHITE, after=8)
        for ins in data.get("insights", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(ins)
            r.font.name = "Arial"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(200, 200, 200)
        if data.get("narrative"):
            add(data["narrative"], size=9, color=GREY, italic=True, after=6)
        if data.get("recommendation"):
            add("STRATEGIC RECOMMENDATION", bold=True, size=7, color=ORANGE, after=3)
            add(data["recommendation"], size=9, color=RGBColor(210, 210, 210), after=10)
        rule()

    add("Confidential  ·  EXD @ Performics  ·  Publicis Groupe", size=7, color=DIM)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# PASSWORD GATE
# ─────────────────────────────────────────────────────────────────────────────

def check_password() -> bool:
    pwd = None
    try:
        pwd = st.secrets["APP_PASSWORD"]
    except Exception:
        pwd = os.environ.get("APP_PASSWORD", "")
    if not pwd:
        return True
    if st.session_state.get("auth"):
        return True

    st.markdown("""
    <div style='text-align:center; padding:4rem 0 2rem;'>
      <div style='font-size:11px; font-weight:700; letter-spacing:0.2em; color:#ff6b2b; text-transform:uppercase; margin-bottom:0.8rem;'>
        Performics · EXD
      </div>
      <div style='font-size:1.7rem; font-weight:700; color:#f0f0f0; margin-bottom:0.25rem;'>
        Insights Generator
      </div>
      <div style='font-size:13px; color:#444; margin-bottom:2.5rem;'>Internal tool — authorised access only</div>
    </div>
    """, unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1.6, 1])
    with col2:
        entered = st.text_input("", type="password", key="pw", placeholder="Enter password")
        if st.button("Access →", use_container_width=True):
            if entered == pwd:
                st.session_state["auth"] = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    return False


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    if not check_password():
        return

    # ── Header ───────────────────────────────────────────────────────────────
    st.markdown("# ⚡ EXD Insights Generator")
    st.markdown("AI-grounded pitch intelligence across all four verticals — in minutes.")
    st.divider()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — Client Brief
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown(
        '<div class="section-hdr">'
        '<div class="section-num">1</div>'
        '<div><div class="section-ttl">Client Brief</div>'
        '<div class="section-sub">Tell us about the pitch</div></div>'
        '</div>',
        unsafe_allow_html=True
    )

    c1, c2 = st.columns(2)
    with c1:
        client_name = st.text_input("Client / Brand", placeholder="e.g. Etihad Airways")
        industry    = st.text_input("Industry", placeholder="e.g. Aviation, Retail Banking")
    with c2:
        market  = st.selectbox("Market / Region", MARKETS, index=0)
        website = st.text_input("Client Website", placeholder="e.g. etihad.com")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Verticals
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown(
        '<div class="section-hdr">'
        '<div class="section-num">2</div>'
        '<div><div class="section-ttl">Verticals</div>'
        '<div class="section-sub">Select which to generate insights for</div></div>'
        '</div>',
        unsafe_allow_html=True
    )

    if "sel_verticals" not in st.session_state:
        st.session_state["sel_verticals"] = list(VERTICALS.keys())

    vert_cols = st.columns(4)
    for i, (vname, vmeta) in enumerate(VERTICALS.items()):
        with vert_cols[i]:
            is_sel = vname in st.session_state["sel_verticals"]
            card_cls = "vert-card selected" if is_sel else "vert-card"
            st.markdown(f"""
            <div class="{card_cls}">
              <span class="vert-icon">{vmeta['icon']}</span>
              <span class="vert-name">{vname}</span>
            </div>""", unsafe_allow_html=True)
            btn_label = "✓ Selected" if is_sel else "Select"
            if st.button(btn_label, key=f"vert_{vname}", use_container_width=True):
                if is_sel:
                    st.session_state["sel_verticals"].remove(vname)
                else:
                    st.session_state["sel_verticals"].append(vname)
                st.rerun()

    selected_verticals = st.session_state["sel_verticals"]
    if not selected_verticals:
        st.caption("⚠️ Select at least one vertical.")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Output Options
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown(
        '<div class="section-hdr">'
        '<div class="section-num">3</div>'
        '<div><div class="section-ttl">Output Options</div>'
        '<div class="section-sub">Customise the report format</div></div>'
        '</div>',
        unsafe_allow_html=True
    )

    # Stats format — card picker
    if "stat_format" not in st.session_state:
        st.session_state["stat_format"] = "headline"

    stat_opts = [
        ("headline", "Headline Stats", "6 punchy data points with source citations"),
        ("narrative", "Narrative Paragraph", "A flowing 220-word market story with embedded stats"),
    ]

    fmt_cols = st.columns(2)
    for i, (val, title, sub) in enumerate(stat_opts):
        with fmt_cols[i]:
            is_sel = st.session_state["stat_format"] == val
            card_cls = "opt-card selected" if is_sel else "opt-card"
            dot_cls  = "opt-dot selected" if is_sel else "opt-dot"
            st.markdown(f"""
            <div class="{card_cls}">
              <div class="{dot_cls}"></div>
              <div>
                <div class="opt-card-title">{title}</div>
                <div class="opt-card-sub">{sub}</div>
              </div>
            </div>""", unsafe_allow_html=True)
            if st.button("Select" if not is_sel else "✓ Selected", key=f"fmt_{val}", use_container_width=True):
                st.session_state["stat_format"] = val
                st.rerun()

    # Narrative toggle — card style
    if "incl_narrative" not in st.session_state:
        st.session_state["incl_narrative"] = False

    incl = st.session_state["incl_narrative"]
    nar_cls = "opt-card selected" if incl else "opt-card"
    dot_cls = "opt-dot selected" if incl else "opt-dot"
    st.markdown(f"""
    <div class="{nar_cls}" style="margin-top:0.5rem;">
      <div class="{dot_cls}"></div>
      <div>
        <div class="opt-card-title">Narrative paragraph per vertical</div>
        <div class="opt-card-sub">Adds a 120-word strategic story to each vertical section</div>
      </div>
    </div>""", unsafe_allow_html=True)
    nar_btn_label = "✓ Enabled" if incl else "Enable"
    if st.button(nar_btn_label, key="toggle_narrative", use_container_width=True):
        st.session_state["incl_narrative"] = not incl
        st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # GENERATE
    # ══════════════════════════════════════════════════════════════════════════
    st.divider()
    ready = all([client_name.strip(), market, industry.strip(), selected_verticals])

    if not ready:
        st.caption("Fill in all fields and select at least one vertical to continue.")

    gen_btn = st.button(
        "⚡  Generate Insights",
        disabled=not ready,
        use_container_width=True,
    )

    if gen_btn and ready:
        st.session_state.pop("results", None)
        st.session_state.pop("meta", None)
        client  = get_gemini_client()
        results = {"verticals": {}}
        total   = 1 + len(selected_verticals)
        prog    = st.progress(0)
        status  = st.empty()

        # Market stats
        status.markdown(
            f'<div class="step-row"><div class="step-dot"></div>Researching {industry} in {market}…</div>',
            unsafe_allow_html=True
        )
        try:
            raw = call_gemini(client, market_stats_prompt(
                client_name, market, industry, website,
                st.session_state["stat_format"]
            ))
            parsed = parse_json(raw)
            results["stats"] = parsed if parsed else {"narrative": raw}
        except Exception as e:
            results["stats_err"] = str(e)
        prog.progress(1 / total)

        # Verticals
        for i, v in enumerate(selected_verticals):
            status.markdown(
                f'<div class="step-row"><div class="step-dot"></div>Generating {v} insights…</div>',
                unsafe_allow_html=True
            )
            try:
                raw = call_gemini(client, vertical_prompt(
                    client_name, market, industry, website,
                    v, VERTICALS[v]["focus"],
                    st.session_state["incl_narrative"]
                ))
                parsed = parse_json(raw)
                results["verticals"][v] = parsed if parsed else {"raw": raw}
            except Exception as e:
                results["verticals"][v] = {"error": str(e)}
            prog.progress((i + 2) / total)

        prog.empty()
        status.empty()
        st.session_state["results"] = results
        st.session_state["meta"] = {
            "client_name": client_name,
            "market": market,
            "industry": industry,
            "website": website,
            "ts": datetime.now().strftime("%d %B %Y, %H:%M"),
        }
        st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # RESULTS
    # ══════════════════════════════════════════════════════════════════════════
    if "results" not in st.session_state:
        return

    results = st.session_state["results"]
    meta    = st.session_state["meta"]

    st.divider()

    st.markdown(f"""
    <div class="result-header">
      <div style='font-size:10px; font-weight:700; letter-spacing:0.2em; text-transform:uppercase; color:#ff6b2b;'>
        Insights Brief
      </div>
      <h2>{meta['client_name']}</h2>
      <p>{meta['market']} &nbsp;·&nbsp; {meta['industry']} &nbsp;·&nbsp; {meta['ts']}</p>
    </div>
    """, unsafe_allow_html=True)

    # Market stats
    if "stats_err" in results:
        st.error(f"Market stats error: {results['stats_err']}")
    elif "stats" in results:
        st.markdown(
            '<div class="section-hdr" style="margin-top:0;">'
            '<div class="section-num" style="background:#1e1e1e; font-size:16px;">📊</div>'
            '<div class="section-ttl">Market Intelligence</div>'
            '</div>',
            unsafe_allow_html=True
        )
        stats = results["stats"]
        if isinstance(stats, list):
            c1, c2 = st.columns(2)
            for i, s in enumerate(stats):
                if isinstance(s, dict):
                    with (c1 if i % 2 == 0 else c2):
                        st.markdown(f"""
                        <div class="stat-card">
                          <div class="stat-num">{s.get('stat','')}</div>
                          <div class="stat-src">{s.get('source','')}</div>
                          <div class="stat-ctx">{s.get('context','')}</div>
                        </div>""", unsafe_allow_html=True)
        elif isinstance(stats, dict) and "narrative" in stats:
            st.markdown(f"""
            <div class="stat-card">
              <div style='font-size:13.5px; color:#aaa; line-height:1.75;'>{stats['narrative']}</div>
            </div>""", unsafe_allow_html=True)

    # Verticals
    for vertical, data in results.get("verticals", {}).items():
        icon = VERTICALS.get(vertical, {}).get("icon", "")
        st.markdown("<hr/>", unsafe_allow_html=True)
        st.markdown(f'<div class="vert-badge">{icon} {vertical}</div>', unsafe_allow_html=True)

        if "error" in data:
            st.error(f"Error: {data['error']}")
            continue
        if "raw" in data:
            with st.expander("Raw output (JSON parse failed)"):
                st.code(data["raw"], language="text")
            continue

        if data.get("headline"):
            st.markdown(f'<div class="headline-pill">{data["headline"]}</div>', unsafe_allow_html=True)
        for ins in data.get("insights", []):
            st.markdown(f'<div class="res-card">{ins}</div>', unsafe_allow_html=True)
        if data.get("narrative"):
            st.markdown(f"""
            <div class="stat-card" style="margin-top:0.4rem;">
              <div style='font-size:13px; color:#666; line-height:1.72; font-style:italic;'>{data['narrative']}</div>
            </div>""", unsafe_allow_html=True)
        if data.get("recommendation"):
            st.markdown(f"""
            <div class="rec-card">
              <div class="rec-label">Strategic Recommendation</div>
              <div class="rec-text">{data['recommendation']}</div>
            </div>""", unsafe_allow_html=True)

    # Export
    st.markdown("<hr/>", unsafe_allow_html=True)
    st.markdown(
        '<div class="section-hdr" style="margin-top:0;">'
        '<div class="section-num" style="background:#1e1e1e;">↓</div>'
        '<div class="section-ttl">Export</div>'
        '</div>',
        unsafe_allow_html=True
    )

    c1, c2, _ = st.columns([1.3, 1.3, 2])
    with c1:
        docx_bytes = generate_docx(meta["client_name"], meta["market"], meta["industry"], results)
        if docx_bytes:
            st.download_button(
                "⬇ Word Doc", data=docx_bytes,
                file_name=f"EXD_Insights_{meta['client_name'].replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.caption("Install python-docx for Word export")

    with c2:
        lines = [
            f"EXD INSIGHTS — {meta['client_name']} | {meta['market']} | {meta['industry']}",
            f"Generated: {meta['ts']}", "=" * 60, "",
        ]
        if "stats" in results:
            lines += ["MARKET INTELLIGENCE", "-" * 40]
            s = results["stats"]
            if isinstance(s, list):
                for item in s:
                    if isinstance(item, dict):
                        lines.append(f"• {item.get('stat','')} ({item.get('source','')})")
                        lines.append(f"  {item.get('context','')}")
            elif isinstance(s, dict) and "narrative" in s:
                lines.append(s["narrative"])
            lines.append("")
        for v, d in results.get("verticals", {}).items():
            lines += [f"\n{v.upper()}", "-" * 40]
            if d.get("headline"): lines.append(f"Headline: {d['headline']}")
            for ins in d.get("insights", []): lines.append(f"\n• {ins}")
            if d.get("recommendation"): lines.append(f"\nRecommendation: {d['recommendation']}")
            lines.append("")
        st.download_button(
            "⬇ Text File", data="\n".join(lines),
            file_name=f"EXD_Insights_{meta['client_name'].replace(' ','_')}.txt",
            mime="text/plain", use_container_width=True,
        )


if __name__ == "__main__":
    main()
